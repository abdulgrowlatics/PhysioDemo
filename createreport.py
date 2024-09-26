from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
import re
import json
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os
logo_path='headerpic.png'





def create_physiotherapy_report(profile_info, logo_path=None):
    def add_custom_text(paragraph, text, bold=False, color=None, size=12):
        run = paragraph.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run.font.size = Pt(size)

    # Create a new Document
    doc = Document()

    # Add the logo to the header if it exists
    if logo_path and os.path.exists(logo_path):
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.0))  # Adjust the width as needed
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Ensure profile_info is a dictionary
    if not isinstance(profile_info, dict):
        raise ValueError("profile_info must be a dictionary.")

    # Add Report Title
    p = doc.add_paragraph()
    add_custom_text(p, "Physiotherapy Assessment Report", bold=True, color=(0, 0, 0), size=16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Clinic Name at the top if provided
    clinic_name = profile_info.get("Clinic Information", {}).get("Address", "Alex's Physiotherapy Assessment")
    if not clinic_name:  # If clinic_name is empty or None, use the default
        clinic_name = "Alex's Physiotherapy Assessment"
    p = doc.add_paragraph()
    add_custom_text(p, f"{clinic_name}", bold=True, size=14)

    # Add Patient Name and Date of Visit if they are provided
    patient_name = profile_info.get('Patient Name', 'N/A')
    if patient_name:
        p = doc.add_paragraph()
        add_custom_text(p, f"Patient Name: {patient_name}", bold=True, color=(0, 0, 0), size=12)
    
    date_of_visit = profile_info.get('Date of Visit', 'N/A')
    if date_of_visit:
        p = doc.add_paragraph()
        add_custom_text(p, f"Date of Visit: {date_of_visit}", bold=True, color=(0, 0, 0), size=12)

    # Add Subjective Findings if they are provided
    subjective_findings = profile_info.get('Subjective Findings', {})
    if isinstance(subjective_findings, dict) and subjective_findings:
        p = doc.add_paragraph()
        add_custom_text(p, "Subjective Findings", bold=True, size=12)
        for key, value in subjective_findings.items():
            if value:  # Only add if the value is not None or empty
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"{key}: {value}", size=10)

    # Add Objective Findings if they are provided
    objective_findings = profile_info.get('Objective Findings', {})
    if isinstance(objective_findings, dict) and objective_findings:
        p = doc.add_paragraph()
        add_custom_text(p, "Objective Findings", bold=True, size=12)
        for key, value in objective_findings.items():
            if isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    if sub_value:  # Only add if the sub_value is not None or empty
                        p = doc.add_paragraph(style='List Bullet')
                        add_custom_text(p, f"{sub_key}: {sub_value}", size=10)
            elif value:  # Only add if the value is not None or empty
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"{key}: {value}", size=10)

    # Add Assessment if it is provided
    assessment = profile_info.get('Assessment', {})
    if isinstance(assessment, dict) and assessment:
        p = doc.add_paragraph()
        add_custom_text(p, "Assessment", bold=True, size=12)
        for key, value in assessment.items():
            if value:  # Only add if the value is not None or empty
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"{key}: {value}", size=10)

    # Add Plan if it is provided and is a dictionary
    plan = profile_info.get('Plan', {})
    if isinstance(plan, dict) and plan:
        p = doc.add_paragraph()
        add_custom_text(p, "Plan", bold=True, size=12)
        for key, value in plan.items():
            if value:  # Only add if the value is not None or empty
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"{key}: {value}", size=10)

    # Add Next Appointment if it is provided
    next_appointment = profile_info.get('Next Appointment', {})
    if isinstance(next_appointment, dict) and next_appointment:
        p = doc.add_paragraph()
        add_custom_text(p, "Next Appointment", bold=True, size=12)
        date = next_appointment.get('Date')
        if date:
            p = doc.add_paragraph(style='List Bullet')
            add_custom_text(p, f"Date: {date}", size=10)
        
        time = next_appointment.get('Time')
        if time:
            p = doc.add_paragraph(style='List Bullet')
            add_custom_text(p, f"Time: {time}", size=10)

    # Add Instructions to Patient if provided
    instructions_to_patient = profile_info.get('Instructions to Patient')
    if instructions_to_patient:
        p = doc.add_paragraph()
        add_custom_text(p, "Instructions to Patient", bold=True, size=12)
        p = doc.add_paragraph(style='List Bullet')
        add_custom_text(p, instructions_to_patient, size=10)

    # Add Clinic Information if provided
    clinic_info = profile_info.get('Clinic Information', {})
    if isinstance(clinic_info, dict) and clinic_info:
        p = doc.add_paragraph()
        add_custom_text(p, "Clinic Information", bold=True, size=12)
        for key, value in clinic_info.items():
            if value:  # Only add if the value is not None or empty
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"{key}: {value}", size=10)

    # Add Footer Information
    footer_text = f"Thank you for choosing {clinic_name} for your care, {patient_name}. Please feel free to reach out if you have any questions or concerns before your next appointment. Together, we'll work towards getting you back to your regular activities pain-free."
    p = doc.add_paragraph()
    add_custom_text(p, footer_text, size=12)
    
    p = doc.add_paragraph()
    add_custom_text(p, "_" * 80, size=10)  # Divider line
    
    # Clinic Contact Information
    clinic_address = clinic_info.get('Address', '123 Fake Street, Springfield')
    clinic_phone = clinic_info.get('Phone', '(555) 123-4567')
    clinic_email = clinic_info.get('Email', 'contact@alexsphysio.com')
    clinic_website = clinic_info.get('Website', 'www.alexsphysio.com')

    p = doc.add_paragraph()
    add_custom_text(p, "Alex's Physiotherapy Assessment Contact Information", bold=True, size=12)
    
    p = doc.add_paragraph(style='List Bullet')
    add_custom_text(p, f"Address: {clinic_address}", size=10)
    
    p = doc.add_paragraph(style='List Bullet')
    add_custom_text(p, f"Phone: {clinic_phone}", size=10)
    
    p = doc.add_paragraph(style='List Bullet')
    add_custom_text(p, f"Email: {clinic_email}", size=10)
    
    p = doc.add_paragraph(style='List Bullet')
    add_custom_text(p, f"Website: {clinic_website}", size=10)

    return doc

def create_exercise_plan_report(exercise_plan, logo_path=None):
    def add_custom_text(paragraph, text, bold=False, color=None, size=12):
        run = paragraph.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run.font.size = Pt(size)

    # Create a new Document
    doc = Document()

    # Add the logo to the header if it exists
    if logo_path and os.path.exists(logo_path):
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.0))  # Adjust the width as needed
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add Report Title
    p = doc.add_paragraph()
    add_custom_text(p, "Physiotherapy Exercise Plan", bold=True, color=(0, 0, 0), size=16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Ensure exercise_plan is a dictionary
    if not isinstance(exercise_plan, dict):
        raise ValueError("exercise_plan must be a dictionary.")

    # Add Patient Name and Date of Visit if they are provided
    patient_name = exercise_plan.get('Patient Name', 'N/A')
    if patient_name:
        p = doc.add_paragraph()
        add_custom_text(p, f"Patient Name: {patient_name}", bold=True, color=(0, 0, 0), size=12)
    
    date_of_visit = exercise_plan.get('Date of Visit', 'N/A')
    if date_of_visit:
        p = doc.add_paragraph()
        add_custom_text(p, f"Date of Visit: {date_of_visit}", bold=True, color=(0, 0, 0), size=12)

    # Add Exercise Plan details
    exercises = exercise_plan.get('Exercise Plan', [])
    if isinstance(exercises, list) and exercises:
        # Add the "Exercise Plan" title
        p = doc.add_paragraph()
        add_custom_text(p, "Exercise Plan", bold=True, size=12)
        
        # Iterate through the list of exercises
        for exercise in exercises:
            # Ensure each exercise is a dictionary
            if isinstance(exercise, dict):
                exercise_name = exercise.get('Exercise Name', 'N/A')
                frequency_repetitions = exercise.get('Frequency/Repetitions', 'N/A')
                description = exercise.get('Description', 'N/A')

                # Add exercise details to the document
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"Exercise Name: {exercise_name}", bold=True, size=10)
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"Frequency/Repetitions: {frequency_repetitions}", size=10)
                p = doc.add_paragraph(style='List Bullet')
                add_custom_text(p, f"Description: {description}", size=10)
                p = doc.add_paragraph()  # Add a blank line for spacing

    return doc



def save_as_docx(profile_info, logo_path):
    # Generate the Word document
    doc = create_physiotherapy_report(profile_info, logo_path)

    # Create a buffer to hold the formatted document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


